import os
import sys

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("python-docx is not installed. Please install it using pip.")
    sys.exit(1)

def generate_detailed_docx():
    print("Initializing Detailed Word Document generation...")
    doc = Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Document Header Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("SECURE AICTE MEETING & GOVERNANCE PLATFORM")
    title_run.font.size = Pt(22)
    title_run.bold = True
    title_run.font.name = 'Arial'
    title_run.font.color.rgb = RGBColor(15, 23, 42)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Module M6: Audit, Attendance, Reports & Institutional Memory\nExhaustive Specifications: React, Node.js + Express, PostgreSQL Database, and Auth/Webhooks")
    sub_run.font.size = Pt(14)
    sub_run.font.name = 'Arial'
    sub_run.font.color.rgb = RGBColor(79, 70, 229)
    
    doc.add_paragraph("").paragraph_format.space_after = Pt(24)

    # 1. Introduction
    doc.add_heading("1. Introduction & Executive Brief", level=1)
    doc.add_paragraph(
        "Technical education governance in India mandates absolute security, transparency, "
        "and record integrity for regulatory committee meetings. The M6 module handles system audit trails, "
        "join/leave duration auditing, compliance reporting, institutional memory searching, and security notifications. "
        "This specification documents the complete structural implementation details of the React frontend, "
        "Node.js + Express backend, PostgreSQL database, and FTS/pgvector search engines."
    )

    # 2. Tech Stack Implementation Deep Dive
    doc.add_heading("2. Tech Stack Architecture Specifications", level=1)

    # 2.1 React
    doc.add_heading("2.1 React + Vite Client-Side Architecture", level=2)
    doc.add_paragraph(
        "The user interface is built as a single-page application (SPA) using React 19 and Vite. "
        "Styling is managed using Tailwind CSS v4 to deliver a modern, secure government portal design. "
        "Key implementation details include:"
    )
    doc.add_paragraph("• Component Layout: Persistent left sidebar navigation coordinating with tab-based state routers (Dashboard, Attendance, Audit, Reports, Memory, Notifications, Settings).", style='List Bullet')
    doc.add_paragraph("• Data Visualization: Implements Recharts graphs (LineChart for attendance rates, PieChart for warning severity breakdowns) displaying compliance metrics.", style='List Bullet')
    doc.add_paragraph("• Unified State Management: Shared React state coordinates real-time dashboard notification counters and threat mitigations.", style='List Bullet')
    doc.add_paragraph("• Guided Presenter Controller: A presentation guide overlay that allows evaluators to run through the demo sequence.", style='List Bullet')

    # 2.2 Node.js + Express
    doc.add_heading("2.2 Node.js + Express Server-Side Logic", level=2)
    doc.add_paragraph(
        "The backend API is constructed using Node.js and Express. It exposes standard REST endpoints "
        "and communicates with database interfaces. Features of the server-side architecture include:"
    )
    doc.add_paragraph("• REST Endpoints: Exposes routes for audit streams, attendance updates, PDF report compilation, and search queries.", style='List Bullet')
    doc.add_paragraph("• Middleware Filters: Integrates CORS filters and JSON body-parsers to process request inputs.", style='List Bullet')
    doc.add_paragraph("• Audit Logging Hook: Post-actions (like report generations or configurations updates) trigger background logging functions.", style='List Bullet')

    # 2.3 PostgreSQL
    doc.add_heading("2.3 PostgreSQL Relational Database Schema & Persistence Fallback", level=2)
    doc.add_paragraph(
        "The database layer is designed for PostgreSQL. It manages relationships between meetings (Module M2), "
        "users (Module M1), and auditing metrics. To resolve server restart data loss, the backend integrates "
        "an automatic database client fallback system:"
    )
    doc.add_paragraph("• PostgreSQL Connection: Initiates a client pool to connect to local PostgreSQL. Creates tables dynamically if missing.", style='List Bullet')
    doc.add_paragraph("• Persistence Fallback (db.json): If PostgreSQL is unreachable, the system activates a local file-based database store inside the workspace, ensuring changes are saved and persisted across server reboots.", style='List Bullet')
    doc.add_paragraph("• Check Constraints: Restricts statuses to designated pools (e.g., Present, Late, Absent, Left Early; INFO, WARNING, CRITICAL) to prevent invalid logs.", style='List Bullet')

    # 2.4 FTS vs pgvector
    doc.add_heading("2.4 PostgreSQL FTS vs. pgvector Search Engine", level=2)
    doc.add_paragraph(
        "To search historic documents and meeting decisions, the platform supports two indexing methods. "
        "Below is a comparison of their implementation details:"
    )
    
    p = doc.add_paragraph()
    r = p.add_run("• PostgreSQL Full-Text Search (FTS): ")
    r.bold = True
    p.add_run(
        "Uses English dictionary lexeme mapping. It compiles document columns into a single 'search_vector' "
        "(tsvector type) and indexes it using a GIN (Generalized Inverted Index) layout. Relevance scoring is computed "
        "via the ts_rank method, ordering results by keyword match accuracy."
    )
    
    p = doc.add_paragraph()
    r = p.add_run("• PostgreSQL pgvector Integration: ")
    r.bold = True
    p.add_run(
        "Uses machine learning embeddings. It stores document representations as vector columns (e.g. vector(1536) "
        "for OpenAI models) and indexes them using HNSW (Hierarchical Navigable Small World) structures. "
        "Similarity searches use cosine distance operators (<=>) to retrieve context-matching records."
    )

    # 3. Database Dictionary & Schemas
    doc.add_heading("3. Database Schema Dictionary", level=1)
    
    doc.add_heading("3.1 Table: audit_logs", level=2)
    table1 = doc.add_table(rows=10, cols=4)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    table1.style = 'Light Shading Accent 1'
    headers = ['Field Name', 'SQL Type', 'Constraints', 'Description']
    for col_idx, text in enumerate(headers):
        table1.rows[0].cells[col_idx].text = text
    cols_data_1 = [
        ('id', 'UUID', 'PRIMARY KEY', 'Unique key generated by gen_random_uuid()'),
        ('user_id', 'UUID', 'FOREIGN KEY', 'References users(id) in Module M1'),
        ('actor_username', 'VARCHAR(50)', 'NOT NULL', 'Username of the actor cached for performance'),
        ('action', 'VARCHAR(100)', 'NOT NULL', 'Description of action (e.g. Permission Changed)'),
        ('module_name', 'VARCHAR(50)', 'NOT NULL', 'Origin module (e.g. Auth (M1))'),
        ('ip_address', 'INET', 'NULLABLE', 'IPv4 or IPv6 network address of the client device'),
        ('details', 'TEXT', 'NULLABLE', 'Context parameters stored in textual JSON format'),
        ('status', 'VARCHAR(25)', 'NOT NULL', 'Success, Failed, Triggered, or Blocked outcome'),
        ('severity_level', 'VARCHAR(15)', 'CHECK', 'INFO, WARNING, or CRITICAL')
    ]
    for row_idx, data in enumerate(cols_data_1):
        for col_idx, val in enumerate(data):
            table1.rows[row_idx+1].cells[col_idx].text = val

    doc.add_heading("3.2 Table: attendance", level=2)
    table2 = doc.add_table(rows=9, cols=4)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.style = 'Light Shading Accent 1'
    for col_idx, text in enumerate(headers):
        table2.rows[0].cells[col_idx].text = text
    cols_data_2 = [
        ('id', 'UUID', 'PRIMARY KEY', 'Unique record identifier'),
        ('meeting_id', 'UUID', 'FOREIGN KEY', 'References meetings(id) in Module M2'),
        ('participant_name', 'VARCHAR(100)', 'NOT NULL', 'Full name of the governance board member'),
        ('official_role', 'VARCHAR(100)', 'NOT NULL', 'Role (e.g. Chairman, CIO, Advisor)'),
        ('join_time', 'TIMESTAMPTZ', 'NULLABLE', 'Timestamp when participant entered the meeting'),
        ('leave_time', 'TIMESTAMPTZ', 'NULLABLE', 'Timestamp when participant exited the meeting'),
        ('duration_minutes', 'INTEGER', 'DEFAULT 0', 'Computed active session time in minutes'),
        ('attendance_status', 'VARCHAR(25)', 'CHECK', 'Present, Late, Absent, or Left Early status')
    ]
    for row_idx, data in enumerate(cols_data_2):
        for col_idx, val in enumerate(data):
            table2.rows[row_idx+1].cells[col_idx].text = val

    # 4. REST API SPECIFICATIONS
    doc.add_heading("4. REST API Endpoint Catalog & JSON Payloads", level=1)
    
    doc.add_heading("4.1 GET /api/audit-logs", level=2)
    doc.add_paragraph("Queries system activity logs. Enforces M1 authentication checks.")
    p = doc.add_paragraph()
    r = p.add_run("• Sample Response Payload:\n")
    r.bold = True
    p.add_run(
        '[\n'
        '  {\n'
        '    "id": "log-102",\n'
        '    "timestamp": "2026-08-08 13:58:10",\n'
        '    "user": "sys_monitor",\n'
        '    "action": "Security Warning",\n'
        '    "module": "Audit (M6)",\n'
        '    "ip": "10.0.4.12",\n'
        '    "status": "Triggered",\n'
        '    "severity": "CRITICAL",\n'
        '    "details": "Multiple failed login attempts detected on admin profile."\n'
        '  }\n'
        ']'
    )
    
    doc.add_heading("4.2 POST /api/reports/generate", level=2)
    doc.add_paragraph("Compiles compliance reports and commits hashes to the blockchain.")
    p = doc.add_paragraph()
    r = p.add_run("• Sample Request Body:\n")
    r.bold = True
    p.add_run(
        '{\n'
        '  "type": "security",\n'
        '  "meetingId": "meet-001",\n'
        '  "dateRange": {\n'
        '    "start": "2026-08-01",\n'
        '    "end": "2026-08-08"\n'
        '  }\n'
        '}'
    )

    doc.add_heading("4.3 GET /api/memory/search", level=2)
    doc.add_paragraph("Searches institutional memory records using text matching. Enforces SQL role-based restriction filters.")
    p = doc.add_paragraph()
    r = p.add_run("• Sample Response Payload:\n")
    r.bold = True
    p.add_run(
        '[\n'
        '  {\n'
        '    "id": "mem-001",\n'
        '    "title": "AICTE Review Meeting - Budget Allocations Q3 2026",\n'
        '    "category": "Meetings",\n'
        '    "date": "2026-08-05",\n'
        '    "relevance": 98,\n'
        '    "details": {\n'
        '      "summary": "AICTE quarterly budget allocation approval. Allocated INR 12.5 Crores.",\n'
        '      "decision": "Approved funding increase of 15% for innovation cell labs.",\n'
        '      "blockchainHash": "0x892e8bf723da890bf2a3e9c8821a9980d28711e9a2bc91e772153c3d2890fb91",\n'
        '      "authorizedRoles": ["Admin", "Chairman", "CIO"]\n'
        '    }\n'
        '  }\n'
        ']'
    )

    doc.add_heading("4.4 POST /api/webhooks/meeting-event", level=2)
    doc.add_paragraph("Webhooks endpoint called by Module M2 Meetings Jitsi server to log participant arrival/exit logs.")
    p = doc.add_paragraph()
    r = p.add_run("• Sample Request Body:\n")
    r.bold = True
    p.add_run(
        '{\n'
        '  "meetingId": "meet-001",\n'
        '  "participant": {\n'
        '    "name": "Dr. Abhay Jere",\n'
        '    "role": "Chief Innovation Officer",\n'
        '    "joinTime": "2026-08-08T10:05:00Z",\n'
        '    "leaveTime": "2026-08-08T11:28:00Z",\n'
        '    "duration": "83",\n'
        '    "status": "Late"\n'
        '  }\n'
        '}'
    )

    # 5. SOP
    doc.add_heading("5. Security SOC Threat Mitigation & M1 RBAC Middleware", level=1)
    doc.add_paragraph("M6 enforces security policies through automated quarantines and M1 identity parsing:")
    doc.add_paragraph("1. M1 Authentication parsing: Middlewares check for 'Authorization' bearer headers (e.g. 'Bearer admin_aicte:Admin'), decoding and binding credentials to requests.", style='List Number')
    doc.add_paragraph("2. Route Guards: Restricted endpoints (like SOC resolve or full system logs) enforce strict role matches, dropping unauthorized requests with 403 Forbidden.", style='List Number')
    doc.add_paragraph("3. IP Quarantine: Blocks client IPs if 5 failed login attempts occur in 2 minutes.", style='List Number')

    # 6. Setup
    doc.add_heading("6. Deployment & Startup Handbook", level=1)
    doc.add_paragraph("To run the prototype, configure the environment and run the startup script:")
    doc.add_paragraph("• Project Location: C:\\Users\\malay\\.gemini\\antigravity\\scratch\\aicte-m6-governance", style='List Bullet')
    doc.add_paragraph("• Port Assignments: Frontend (5173), Express Backend (5000)", style='List Bullet')
    doc.add_paragraph("• Run Command: npm start (runs both systems concurrently)", style='List Bullet')

    doc.save("AICTE_M6_Technical_Documentation_Detailed.docx")
    print("Exhaustive detailed Word Document updated successfully as 'AICTE_M6_Technical_Documentation_Detailed.docx'")

if __name__ == "__main__":
    generate_detailed_docx()
